"""
flowworkshop · services/checklist_export_service.py

Erzeugt aus einem KOM-Checklisten-Template (models/checklist_template.py) eine
ausfuellbare Pruefcheckliste in den Formaten DOCX, XLSX und PDF.

Layout (gespiegelt zum audit_designer-Checklistendesigner):
  - Deckblatt/Kopfblock aus ``ChecklistTemplate.properties_json`` (Audit code,
    CCI/Programme, prepared/reviewed Name+Date …)
  - Inhalt kapitelweise (HEADING gruppiert) als Tabelle mit den Spalten
    Nr · Legal reference · Frage · Antwort · Bemerkung · Belege.
  - HINT-Knoten erscheinen als eingerueckte Hinweis-Unterzeile UNTER der Frage.
  - DECISION-Knoten mit JA/NEIN-Zweigen werden eingerueckt dargestellt.
  - Die Antwort-Spalte rendert die Optionen des zugewiesenen Antwortsets als
    ankreuzbare Felder (z.B. ☐ Ja  ☐ Nein  ☐ Entfaellt); bei Betrag/Datum/
    Freitext ein Leerfeld.

Zwei Modi:
  - ``blank``  : leere, ankreuzbare Felder (Standard) — die Checkliste wird
    erst in der Pruefung ausgefuellt.
  - ``filled`` : vorbefuellt, falls Antwortdaten vorliegen. Da ein Template in
    der Regel keine Antworten enthaelt, ist das Ergebnis meist identisch zu
    ``blank``.

Im DOCX werden fuer die ausfuellbaren Antwort-, Bemerkungs- und Kopf-Felder
echte Word-Inhaltssteuerelemente (w:sdt Plain-Text) verwendet, damit sie nach
dem Export in Word direkt bearbeitbar sind.

Eigenstaendiger Helfer — bewusst getrennt vom projektgebundenen
services/excel_export.py. Liest Modelle aus models/checklist_template.py.
"""
from __future__ import annotations

import io
import logging
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Pt, RGBColor

# w14-Namespace (Word 2010) registrieren, damit qn('w14:checkbox') aufloesbar ist.
# Wird fuer echte, anklickbare Word-Checkbox-Inhaltssteuerelemente benoetigt.
nsmap.setdefault("w14", "http://schemas.microsoft.com/office/word/2010/wordml")

from openpyxl import Workbook  # noqa: E402
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side  # noqa: E402
from openpyxl.utils import get_column_letter  # noqa: E402

from models.checklist_template import (  # noqa: E402
    ChecklistTemplate,
    ChecklistTemplateNode,
    ChecklistAnswerSet,
)

log = logging.getLogger(__name__)


# ── Farben (an audit_designer angelehnt, aber im Workshop-Emerald) ────────────
HEADER_BG = "0E7C5C"      # Workshop-Emerald
HEADING_BG = "D9EAE3"     # heller Emerald-Ton fuer Kapitelzeilen
DECISION_BG = "E8EEF7"    # helles Blau fuer Entscheidungen
BRANCH_JA_BG = "E8F5E9"   # helles Gruen fuer JA-Zweig
BRANCH_NEIN_BG = "FFF3E0"  # helles Orange fuer NEIN-Zweig
HINT_BG = "FFFBEB"        # helles Gelb fuer Hinweise
GRAY_BG = "F1F5F9"        # mildes Grau

# Leerkaestchen / Ankreuz-Symbole
BOX_EMPTY = "☐"   # ☐
BOX_CHECK = "☒"   # ☒

# Knotentyp-Bezeichnungen (Anzeige)
NODE_TYPE_LABELS = {
    "HEADING": "Kapitel",
    "QUESTION": "Frage",
    "DECISION": "Entscheidung",
    "HINT": "Hinweis",
}

# Standard-Optionen je answer_type, falls KEIN Antwortset zugewiesen ist.
_DEFAULT_OPTIONS_BY_TYPE = {
    "BOOLEAN": ["Ja", "Nein", "Teilweise", "Entfaellt"],
    "BOOLEAN_JN": ["Ja", "Nein"],
    "CUSTOM_ENUM": [],   # ohne Set keine festen Optionen
}

# eingabetyp (QChess FRAGENTYPID): 0=Auswahl,1=Freitext,2=Betrag,4=Datum
_EINGABETYP_FREITEXT = 1
_EINGABETYP_BETRAG = 2
_EINGABETYP_DATUM = 4


# ── Datenaufbereitung: Baum + Kopfblock ───────────────────────────────────────

def _kopfblock_items(props: dict | None) -> list[tuple[str, str]]:
    """Baut aus ``properties_json`` eine geordnete Liste (Label, Wert) fuer den
    Kopfblock.

    Akzeptiert sowohl die KOM-typischen englischen Schluessel (audit_code, cci,
    programme, prepared_by/prepared_date, reviewed_by/reviewed_date) als auch
    freie Zusatzfelder. Unbekannte Schluessel werden generisch angehaengt, damit
    nichts verloren geht."""
    props = props or {}

    # Bekannte Schluessel in fester Reihenfolge mit deutschsprachigem Label.
    known: list[tuple[str, list[str]]] = [
        ("Audit code", ["audit_code", "auditCode", "audit_no", "aktenzeichen"]),
        ("CCI", ["cci", "CCI"]),
        ("Programm", ["programme", "program", "programm"]),
        ("Prioritaetsachse", ["priority_axis", "prioritaetsachse"]),
        ("Fonds", ["fund", "fonds"]),
        ("Geschaeftsjahr", ["accounting_year", "geschaeftsjahr"]),
        ("Erstellt von", ["prepared_by", "preparedBy", "prepared", "erstellt_von"]),
        ("Erstellt am", ["prepared_date", "preparedDate", "erstellt_am"]),
        ("Geprueft von", ["reviewed_by", "reviewedBy", "reviewed", "geprueft_von"]),
        ("Geprueft am", ["reviewed_date", "reviewedDate", "geprueft_am"]),
    ]

    used_keys: set[str] = set()
    items: list[tuple[str, str]] = []
    for label, candidates in known:
        for key in candidates:
            if key in props and props.get(key) not in (None, ""):
                items.append((label, str(props.get(key))))
                used_keys.update(candidates)
                break
        else:
            # Auch leere bekannte Felder als ausfuellbare Zeile vorsehen.
            used_keys.update(candidates)

    # Restliche, unbekannte Schluessel generisch anhaengen.
    for key, value in props.items():
        if key in used_keys:
            continue
        if value in (None, ""):
            continue
        label = str(key).replace("_", " ").strip().capitalize()
        items.append((label, str(value)))

    return items


def _build_node_tree(
    nodes: list[ChecklistTemplateNode],
) -> tuple[dict[str, dict], list[str]]:
    """Baut aus der flachen Knotenliste eine children-Map + Wurzel-Reihenfolge.

    Rueckgabe:
      ``by_id``   : {node_id: {"node": ChecklistTemplateNode, "children": [ids]}}
      ``roots``   : sortierte Liste der Wurzel-Knoten-IDs

    HINT-Knoten werden — wie im Frontend-Baum — ans Ende ihrer Geschwister-Ebene
    sortiert, damit ein Hinweis unterhalb der zugehoerigen Frage erscheint."""
    known_ids = {n.id for n in nodes}
    by_id: dict[str, dict] = {n.id: {"node": n, "children": []} for n in nodes}

    children_map: dict[str | None, list[ChecklistTemplateNode]] = {}
    for node in nodes:
        parent = node.parent_id if node.parent_id in known_ids else None
        children_map.setdefault(parent, []).append(node)

    def _sort_key(n: ChecklistTemplateNode) -> tuple:
        is_hint = 1 if (n.node_type or "").upper() == "HINT" else 0
        return (n.sort_order or 0, is_hint, n.created_at or datetime.min)

    for parent_id, kids in children_map.items():
        ordered = sorted(kids, key=_sort_key)
        if parent_id is not None and parent_id in by_id:
            by_id[parent_id]["children"] = [k.id for k in ordered]

    roots = [n.id for n in sorted(children_map.get(None, []), key=_sort_key)]
    return by_id, roots


def _answer_options_for(
    node: ChecklistTemplateNode, answer_sets: dict[str, ChecklistAnswerSet],
) -> list[str]:
    """Ermittelt die ankreuzbaren Optionen-Labels fuer eine Frage/Entscheidung.

    Reihenfolge der Aufloesung:
      1. Zugewiesenes Antwortset (answer_set_id) → dessen Optionen.
      2. Sonst Standard-Optionen aus dem answer_type (BOOLEAN/BOOLEAN_JN).
    Bei Freitext/Betrag/Datum (eingabetyp) werden KEINE Ankreuz-Optionen
    geliefert — dort entsteht ein Leerfeld (siehe _render_*)."""
    aset = answer_sets.get(node.answer_set_id) if node.answer_set_id else None
    if aset and aset.options:
        opts = sorted(aset.options, key=lambda o: (o.sort_order or 0, o.name or ""))
        return [o.name for o in opts if o.name]

    atype = (node.answer_type or "").upper()
    return list(_DEFAULT_OPTIONS_BY_TYPE.get(atype, []))


def _is_freitext_like(node: ChecklistTemplateNode) -> bool:
    """True, wenn die Antwort ein Leerfeld (Freitext/Betrag/Datum) statt
    Ankreuz-Optionen sein soll."""
    if node.eingabetyp in (_EINGABETYP_FREITEXT, _EINGABETYP_BETRAG, _EINGABETYP_DATUM):
        return True
    atype = (node.answer_type or "").upper()
    return atype in ("TEXT", "CURRENCY", "DATE")


def _freitext_hint(node: ChecklistTemplateNode) -> str:
    """Beschriftung des Leerfelds je nach eingabetyp/answer_type."""
    if node.eingabetyp == _EINGABETYP_BETRAG or (node.answer_type or "").upper() == "CURRENCY":
        return "Betrag (EUR): __________"
    if node.eingabetyp == _EINGABETYP_DATUM or (node.answer_type or "").upper() == "DATE":
        return "Datum: __.__.____"
    return "________________________"


def _relevant_documents(node: ChecklistTemplateNode) -> str:
    """Serialisiert relevant_documents_json zu einem lesbaren Mehrzeiler."""
    raw = node.relevant_documents_json
    if not raw:
        return ""
    if isinstance(raw, str):
        return raw
    if isinstance(raw, list):
        parts: list[str] = []
        for item in raw:
            if isinstance(item, dict):
                parts.append(str(item.get("name") or item.get("title") or item.get("label") or item))
            else:
                parts.append(str(item))
        return "\n".join(p for p in parts if p)
    if isinstance(raw, dict):
        return "\n".join(f"{k}: {v}" for k, v in raw.items())
    return str(raw)


def _node_title(node: ChecklistTemplateNode) -> str:
    """Anzeigetext eines Knotens (bevorzugt geprueft/uebersetzt, sonst title)."""
    return (
        node.review_text_de
        or node.translated_text_de
        or node.title
        or node.source_text_en
        or ""
    ).strip()


class _RowCollector:
    """Sammelt aus dem Knotenbaum eine flache, kapitelweise gruppierte
    Zeilenliste fuer die drei Renderer (DOCX/XLSX/PDF).

    Jede Zeile ist ein Dict:
      nr            : hierarchische Nummer (z.B. "1.2.3")
      legal         : Rechtsgrundlage (legal_reference)
      question      : Fragetext (mit JA/NEIN-Aussage bei DECISION)
      options       : Liste ankreuzbarer Option-Labels (oder [])
      freitext_hint : Beschriftung Leerfeld (falls Freitext/Betrag/Datum)
      remark        : Bemerkung (public_remark; im filled-Modus vorbefuellt)
      documents     : Belegverweise (mehrzeilig)
      row_type      : "heading" | "question" | "decision" | "branch" | "hint"
      indent        : Einrueckungstiefe (0 = oberste Ebene)
      branch        : "" | "JA" | "NEIN"
    """

    def __init__(
        self,
        by_id: dict[str, dict],
        answer_sets: dict[str, ChecklistAnswerSet],
        mode: str,
    ):
        self.by_id = by_id
        self.answer_sets = answer_sets
        self.mode = mode
        # chapters: Liste (Kapiteltitel, [Zeilen])
        self.chapters: list[tuple[str, list[dict]]] = []
        self._current_title = ""
        self._current_rows: list[dict] = []

    def _flush(self) -> None:
        if self._current_rows or self._current_title:
            self.chapters.append((self._current_title, self._current_rows))
        self._current_title = ""
        self._current_rows = []

    def collect(self, roots: list[str]) -> list[tuple[str, list[dict]]]:
        # Zaehler je Hierarchieebene fuer die Nummerierung.
        for root_id in roots:
            self._walk(root_id, level=0, indent=0, prefix=[])
        self._flush()
        return self.chapters

    def _walk(self, node_id: str, level: int, indent: int, prefix: list[int]) -> None:
        entry = self.by_id.get(node_id)
        if not entry:
            return
        node: ChecklistTemplateNode = entry["node"]
        node_type = (node.node_type or "QUESTION").upper()
        title = _node_title(node)
        children = entry["children"]

        if node_type == "HEADING":
            # Oberste Kapitel beenden den vorherigen Block.
            if level == 0:
                self._flush()
                self._current_title = title or "Kapitel"
                child_prefix: list[int] = []
            else:
                # Unter-Heading als eigene Zeile.
                self._current_rows.append({
                    "nr": "", "legal": "", "question": title,
                    "options": [], "freitext_hint": "", "remark": "",
                    "documents": "", "row_type": "heading", "indent": indent,
                    "branch": "",
                })
                child_prefix = list(prefix)
            counter = 0
            for child_id in children:
                counter += 1
                self._walk(child_id, level + 1, indent, child_prefix + [counter])
            return

        if node_type == "DECISION":
            nr = ".".join(str(p) for p in prefix) if prefix else ""
            ja_label = (node.ja_label or "").strip()
            nein_label = (node.nein_label or "").strip()
            self._current_rows.append({
                "nr": nr, "legal": (node.legal_reference or "").strip(),
                "question": title, "options": self._options(node),
                "freitext_hint": self._freitext(node),
                "remark": self._remark(node), "documents": _relevant_documents(node),
                "row_type": "decision", "indent": indent, "branch": "",
            })
            # Kinder nach Zweig gruppieren.
            ja_kids, nein_kids, other_kids = self._split_branches(children)
            sub_counter = 0
            for child_id in other_kids:
                sub_counter += 1
                self._walk(child_id, level + 1, indent + 1, prefix + [sub_counter])
            if ja_kids or ja_label:
                self._current_rows.append({
                    "nr": "", "legal": "",
                    "question": ja_label or f"Wenn {nr} = Ja:",
                    "options": [], "freitext_hint": "", "remark": "",
                    "documents": "", "row_type": "branch", "indent": indent,
                    "branch": "JA",
                })
                for child_id in ja_kids:
                    sub_counter += 1
                    self._walk(child_id, level + 1, indent + 1, prefix + [sub_counter])
            if nein_kids or nein_label:
                self._current_rows.append({
                    "nr": "", "legal": "",
                    "question": nein_label or f"Wenn {nr} = Nein:",
                    "options": [], "freitext_hint": "", "remark": "",
                    "documents": "", "row_type": "branch", "indent": indent,
                    "branch": "NEIN",
                })
                for child_id in nein_kids:
                    sub_counter += 1
                    self._walk(child_id, level + 1, indent + 1, prefix + [sub_counter])
            return

        if node_type == "HINT":
            self._current_rows.append({
                "nr": "", "legal": "", "question": title,
                "options": [], "freitext_hint": "", "remark": "",
                "documents": "", "row_type": "hint", "indent": indent, "branch": "",
            })
            # Hinweise haben i.d.R. keine Kinder; falls doch, anhaengen.
            for child_id in children:
                self._walk(child_id, level + 1, indent, prefix)
            return

        # QUESTION
        nr = ".".join(str(p) for p in prefix) if prefix else ""
        self._current_rows.append({
            "nr": nr, "legal": (node.legal_reference or "").strip(),
            "question": title, "options": self._options(node),
            "freitext_hint": self._freitext(node),
            "remark": self._remark(node), "documents": _relevant_documents(node),
            "row_type": "question", "indent": indent, "branch": "",
        })
        ja_kids, nein_kids, other_kids = self._split_branches(children)
        sub_counter = 0
        for child_id in other_kids + ja_kids + nein_kids:
            sub_counter += 1
            self._walk(child_id, level + 1, indent + 1, prefix + [sub_counter])

    def _split_branches(self, children: list[str]) -> tuple[list[str], list[str], list[str]]:
        ja_kids, nein_kids, other_kids = [], [], []
        for child_id in children:
            entry = self.by_id.get(child_id)
            if not entry:
                continue
            branch = (entry["node"].branch or "").upper()
            if branch == "JA":
                ja_kids.append(child_id)
            elif branch == "NEIN":
                nein_kids.append(child_id)
            else:
                other_kids.append(child_id)
        return ja_kids, nein_kids, other_kids

    def _options(self, node: ChecklistTemplateNode) -> list[str]:
        if _is_freitext_like(node):
            return []
        return _answer_options_for(node, self.answer_sets)

    def _freitext(self, node: ChecklistTemplateNode) -> str:
        return _freitext_hint(node) if _is_freitext_like(node) else ""

    def _remark(self, node: ChecklistTemplateNode) -> str:
        """Bemerkung — im filled-Modus die hinterlegte public_remark, sonst leer."""
        if self.mode == "filled":
            return (node.public_remark or "").strip()
        return ""


def _prepare(
    template: ChecklistTemplate,
    nodes: list[ChecklistTemplateNode],
    answer_sets: list[ChecklistAnswerSet],
    mode: str,
) -> tuple[list[tuple[str, str]], list[tuple[str, list[dict]]]]:
    """Gemeinsame Vorbereitung fuer alle drei Formate.

    Rueckgabe: (kopfblock_items, chapters)."""
    aset_map = {a.id: a for a in answer_sets}
    by_id, roots = _build_node_tree(nodes)
    collector = _RowCollector(by_id, aset_map, mode)
    chapters = collector.collect(roots)
    kopf = _kopfblock_items(template.properties_json)
    return kopf, chapters


# ── DOCX-Helfer: Zellenformatierung + echte Word-Formularfelder ────────────────

def _set_cell_shading(cell, color: str) -> None:
    """Setzt die Hintergrundfarbe einer Tabellenzelle (w:shd)."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell) -> None:
    """Setzt duenne schwarze Rahmen an allen vier Seiten einer Zelle."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        borders.append(b)
    tcPr.append(borders)


def _add_sdt_text(paragraph, placeholder: str = "", italic: bool = False) -> None:
    """Fuegt einem Paragraphen ein echtes Word-Inhaltssteuerelement (Plain-Text
    w:sdt) hinzu — nach dem Export in Word direkt ausfuellbar/anklickbar.

    Erzeugt manuell die OOXML-Struktur:
        <w:sdt>
          <w:sdtPr><w:alias/><w:tag/><w:id/><w:text/></w:sdtPr>
          <w:sdtContent><w:r><w:t>placeholder</w:t></w:r></w:sdtContent>
        </w:sdt>
    Der Platzhaltertext bleibt erhalten und wird beim ersten Klick in Word
    ueberschrieben."""
    sdt = OxmlElement("w:sdt")

    sdt_pr = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), "Eingabe")
    sdt_pr.append(alias)
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "eingabe")
    sdt_pr.append(tag)
    sdt_id = OxmlElement("w:id")
    # Eindeutige, stabile ID je Element (positiver 31-Bit-Wert).
    sdt_id.set(qn("w:val"), str(abs(hash((id(paragraph), placeholder))) % 2_000_000_000))
    sdt_pr.append(sdt_id)
    sdt_pr.append(OxmlElement("w:text"))
    sdt.append(sdt_pr)

    sdt_content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "808080" if (italic or not placeholder) else "000000")
    rpr.append(color)
    if italic:
        rpr.append(OxmlElement("w:i"))
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = placeholder or " "
    run.append(text_el)
    sdt_content.append(run)
    sdt.append(sdt_content)

    paragraph._p.append(sdt)


# Globaler, monoton steigender Zaehler fuer eindeutige w:sdt-Checkbox-IDs.
_checkbox_id_counter = 0


def _add_checkbox_option(paragraph, label: str, checked: bool = False) -> None:
    """Fuegt dem Paragraphen ein echtes, in Word anklickbares Checkbox-
    Inhaltssteuerelement (``w14:checkbox``) gefolgt von der Beschriftung hinzu.

    Die Box laesst sich nach dem Export in Word per Klick zwischen ☐ und ☒
    umschalten. Erzeugt folgende OOXML-Struktur:

        <w:sdt>
          <w:sdtPr>
            <w14:checkbox>
              <w14:checked w14:val="0|1"/>
              <w14:checkedState w14:val="2612" w14:font="MS Gothic"/>
              <w14:uncheckedState w14:val="2610" w14:font="MS Gothic"/>
            </w14:checkbox>
          </w:sdtPr>
          <w:sdtContent><w:r>…☐/☒…</w:r></w:sdtContent>
        </w:sdt>

    Danach folgt im selben Paragraphen ein normaler Run " {label}"."""
    global _checkbox_id_counter
    _checkbox_id_counter += 1

    sdt = OxmlElement("w:sdt")

    sdt_pr = OxmlElement("w:sdtPr")
    sdt_id = OxmlElement("w:id")
    sdt_id.set(qn("w:val"), str(_checkbox_id_counter))
    sdt_pr.append(sdt_id)

    checkbox = OxmlElement("w14:checkbox")
    checked_el = OxmlElement("w14:checked")
    checked_el.set(qn("w14:val"), "1" if checked else "0")
    checkbox.append(checked_el)
    checked_state = OxmlElement("w14:checkedState")
    checked_state.set(qn("w14:val"), "2612")
    checked_state.set(qn("w14:font"), "MS Gothic")
    checkbox.append(checked_state)
    unchecked_state = OxmlElement("w14:uncheckedState")
    unchecked_state.set(qn("w14:val"), "2610")
    unchecked_state.set(qn("w14:font"), "MS Gothic")
    checkbox.append(unchecked_state)
    sdt_pr.append(checkbox)
    sdt.append(sdt_pr)

    sdt_content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "MS Gothic")
    rfonts.set(qn("w:hAnsi"), "MS Gothic")
    rfonts.set(qn("w:eastAsia"), "MS Gothic")
    rpr.append(rfonts)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = BOX_CHECK if checked else BOX_EMPTY
    run.append(text_el)
    sdt_content.append(run)
    sdt.append(sdt_content)

    paragraph._p.append(sdt)

    # Beschriftung als normaler Run im selben Paragraphen.
    label_run = paragraph.add_run(f" {label}")
    label_run.font.size = Pt(9)


def _docx_answer_cell(cell, row: dict, mode: str) -> None:
    """Fuellt die Antwort-Zelle: anklickbare Checkbox-Optionen ODER ausfuellbares
    Leerfeld (echte Word-Inhaltssteuerelemente)."""
    cell.paragraphs[0].clear()
    options = row.get("options") or []
    freitext = row.get("freitext_hint") or ""

    if options:
        # Pro Option eine Zeile mit echter, anklickbarer Word-Checkbox.
        for idx, opt in enumerate(options):
            target = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            _add_checkbox_option(target, opt, checked=False)
    elif freitext:
        p = cell.paragraphs[0]
        run = p.add_run("")
        run.font.size = Pt(9)
        _add_sdt_text(p, placeholder=freitext, italic=True)
    else:
        # generisches ausfuellbares Feld
        _add_sdt_text(cell.paragraphs[0], placeholder="________", italic=True)


# ── DOCX-Export ────────────────────────────────────────────────────────────────

def export_docx(
    template: ChecklistTemplate,
    nodes: list[ChecklistTemplateNode],
    answer_sets: list[ChecklistAnswerSet],
    mode: str = "blank",
) -> bytes:
    """Erzeugt die ausfuellbare Pruefcheckliste als DOCX (Bytes).

    Verwendet echte Word-Inhaltssteuerelemente (w:sdt) fuer die ausfuellbaren
    Antwort-, Bemerkungs- und Kopfblock-Felder."""
    kopf, chapters = _prepare(template, nodes, answer_sets, mode)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ── Deckblatt ──
    header_table = doc.add_table(rows=1, cols=1)
    header_cell = header_table.rows[0].cells[0]
    _set_cell_shading(header_cell, HEADER_BG)
    title_para = header_cell.paragraphs[0]
    title_run = title_para.add_run("Pruefcheckliste")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(255, 255, 255)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(14)
    sub_para = header_cell.add_paragraph()
    sub_run = sub_para.add_run(template.title or "")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(255, 255, 255)
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    if template.description:
        desc = doc.add_paragraph()
        desc.add_run(template.description.strip()).italic = True

    # ── Kopfblock (ausfuellbar) ──
    if kopf:
        info_table = doc.add_table(rows=len(kopf), cols=2)
        info_table.style = "Table Grid"
        for i, (label, value) in enumerate(kopf):
            row = info_table.rows[i]
            row.cells[0].text = label
            if row.cells[0].paragraphs[0].runs:
                row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[0].width = Cm(5)
            _set_cell_shading(row.cells[0], GRAY_BG)
            value_cell = row.cells[1]
            value_cell.paragraphs[0].clear()
            value_cell.width = Cm(12)
            if mode == "filled" and value:
                value_cell.paragraphs[0].add_run(str(value)).font.size = Pt(10)
            else:
                # Leeres, ausfuellbares Kopffeld (vorbelegt mit ggf. vorhandenem Wert).
                _add_sdt_text(value_cell.paragraphs[0], placeholder=str(value) if value else "", italic=not value)

    doc.add_paragraph()
    gen = doc.add_paragraph()
    gen_run = gen.add_run(f"Erzeugt am {datetime.now().strftime('%d.%m.%Y %H:%M')} · Modus: {mode}")
    gen_run.font.size = Pt(8)
    gen_run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # ── Inhalt kapitelweise ──
    COL_WIDTHS = [Cm(1.2), Cm(3.0), Cm(6.5), Cm(3.3), Cm(3.5), Cm(2.5)]
    HEADERS = ["Nr.", "Rechtsgrundlage", "Frage", "Antwort", "Bemerkung", "Belege"]
    INDENT = "    "

    for chapter_title, rows in chapters:
        if chapter_title:
            doc.add_heading(chapter_title, level=1)
        if not rows:
            continue

        table = doc.add_table(rows=1 + len(rows), cols=len(HEADERS))
        table.style = "Table Grid"
        table.autofit = False

        for ci, label in enumerate(HEADERS):
            cell = table.rows[0].cells[ci]
            cell.width = COL_WIDTHS[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, HEADER_BG)
            _set_cell_borders(cell)

        for ri, row in enumerate(rows):
            r = table.rows[ri + 1]
            row_type = row["row_type"]
            indent_prefix = INDENT * row.get("indent", 0)
            for ci in range(len(HEADERS)):
                r.cells[ci].width = COL_WIDTHS[ci]
                _set_cell_borders(r.cells[ci])

            # Heading / Branch / Hint = ueber alle Spalten zusammengefuehrte Zeile
            if row_type in ("heading", "branch", "hint"):
                merged = r.cells[0].merge(r.cells[len(HEADERS) - 1])
                merged.paragraphs[0].clear()
                if row_type == "hint":
                    text = f"{indent_prefix}Hinweis: {row['question']}"
                else:
                    text = f"{indent_prefix}{row['question']}"
                run = merged.paragraphs[0].add_run(text)
                run.font.size = Pt(9)
                run.bold = row_type in ("heading", "branch")
                run.italic = row_type in ("branch", "hint")
                bg = {
                    "heading": HEADING_BG,
                    "hint": HINT_BG,
                    "branch": BRANCH_JA_BG if row.get("branch") == "JA" else BRANCH_NEIN_BG,
                }[row_type]
                _set_cell_shading(merged, bg)
                continue

            # Nr.
            nr_cell = r.cells[0]
            nr_cell.paragraphs[0].clear()
            nr_run = nr_cell.paragraphs[0].add_run(row["nr"])
            nr_run.font.size = Pt(9)
            nr_run.bold = row_type == "decision"
            nr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Rechtsgrundlage
            legal_cell = r.cells[1]
            legal_cell.paragraphs[0].clear()
            legal_cell.paragraphs[0].add_run(row["legal"]).font.size = Pt(8)

            # Frage (+ ggf. nichts; HINT ist eigene Zeile)
            q_cell = r.cells[2]
            q_cell.paragraphs[0].clear()
            q_run = q_cell.paragraphs[0].add_run(f"{indent_prefix}{row['question']}")
            q_run.font.size = Pt(9)
            q_run.bold = row_type == "decision"

            # Antwort (ankreuzbar / Leerfeld)
            _docx_answer_cell(r.cells[3], row, mode)

            # Bemerkung (ausfuellbar)
            b_cell = r.cells[4]
            b_cell.paragraphs[0].clear()
            if mode == "filled" and row.get("remark"):
                b_cell.paragraphs[0].add_run(row["remark"]).font.size = Pt(9)
            else:
                _add_sdt_text(b_cell.paragraphs[0], placeholder="", italic=True)

            # Belege
            d_cell = r.cells[5]
            d_cell.paragraphs[0].clear()
            d_cell.paragraphs[0].add_run(row.get("documents", "")).font.size = Pt(8)

            if row_type == "decision":
                for ci in range(len(HEADERS)):
                    _set_cell_shading(r.cells[ci], DECISION_BG)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── XLSX-Export ────────────────────────────────────────────────────────────────

def export_xlsx(
    template: ChecklistTemplate,
    nodes: list[ChecklistTemplateNode],
    answer_sets: list[ChecklistAnswerSet],
    mode: str = "blank",
) -> bytes:
    """Erzeugt die ausfuellbare Pruefcheckliste als XLSX (Bytes).

    Flache, gut lesbare Tabelle. Header-Styling analog services/excel_export.py
    (Workshop-Emerald, weisse Schrift, AutoFilter, eingefrorene Kopfzeile)."""
    kopf, chapters = _prepare(template, nodes, answer_sets, mode)

    wb = Workbook()
    header_fill = PatternFill("solid", fgColor=HEADER_BG)
    header_font = Font(bold=True, color="FFFFFF", size=11)
    heading_fill = PatternFill("solid", fgColor=HEADING_BG)
    decision_fill = PatternFill("solid", fgColor=DECISION_BG)
    branch_ja_fill = PatternFill("solid", fgColor=BRANCH_JA_BG)
    branch_nein_fill = PatternFill("solid", fgColor=BRANCH_NEIN_BG)
    hint_fill = PatternFill("solid", fgColor=HINT_BG)
    title_font = Font(bold=True, size=14)
    thin = Side(style="thin", color="D0D7DE")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap_top = Alignment(wrap_text=True, vertical="top")

    # ── Sheet 1: Kopfblock ──
    ws_info = wb.active
    ws_info.title = "Kopfblock"
    ws_info.cell(row=1, column=1, value="Pruefcheckliste").font = title_font
    ws_info.cell(row=2, column=1, value=template.title or "").font = Font(bold=True, size=12)
    if template.description:
        ws_info.cell(row=3, column=1, value=template.description.strip()).alignment = wrap_top
    r = 5
    for label, value in kopf:
        ws_info.cell(row=r, column=1, value=label).font = Font(bold=True)
        ws_info.cell(row=r, column=2, value=(value if mode == "filled" else value) or "")
        r += 1
    ws_info.cell(row=r + 1, column=1, value="Erzeugt").font = Font(bold=True)
    ws_info.cell(row=r + 1, column=2, value=datetime.now().strftime("%d.%m.%Y %H:%M"))
    ws_info.cell(row=r + 2, column=1, value="Modus").font = Font(bold=True)
    ws_info.cell(row=r + 2, column=2, value=mode)
    ws_info.column_dimensions["A"].width = 28
    ws_info.column_dimensions["B"].width = 70

    # ── Sheet 2: Checkliste (flach) ──
    ws = wb.create_sheet("Checkliste")
    headers = ["Pfad/Nr.", "Typ", "Frage", "Antwortoptionen", "Bemerkung", "Belege", "Rechtsgrundlage"]
    for ci, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[1].height = 22

    row_idx = 2
    for chapter_title, rows in chapters:
        if chapter_title:
            cell = ws.cell(row=row_idx, column=1, value=chapter_title)
            ws.merge_cells(start_row=row_idx, start_column=1, end_row=row_idx, end_column=len(headers))
            cell.font = Font(bold=True, size=11)
            for ci in range(1, len(headers) + 1):
                ws.cell(row=row_idx, column=ci).fill = heading_fill
            row_idx += 1

        for row in rows:
            row_type = row["row_type"]
            indent = "    " * row.get("indent", 0)
            if row_type == "hint":
                question = f"{indent}Hinweis: {row['question']}"
            else:
                question = f"{indent}{row['question']}"

            if row["options"]:
                ans = "  ".join(f"{BOX_EMPTY} {o}" for o in row["options"])
            elif row.get("freitext_hint"):
                ans = row["freitext_hint"]
            elif row_type in ("question", "decision"):
                ans = "________"
            else:
                ans = ""

            values = [
                row["nr"],
                NODE_TYPE_LABELS.get(
                    {"heading": "HEADING", "question": "QUESTION", "decision": "DECISION",
                     "branch": "HEADING", "hint": "HINT"}[row_type], ""),
                question,
                ans,
                row.get("remark", ""),
                row.get("documents", ""),
                row.get("legal", ""),
            ]
            for ci, val in enumerate(values, start=1):
                cell = ws.cell(row=row_idx, column=ci, value=val)
                cell.border = border
                if ci in (3, 4, 5, 6, 7):
                    cell.alignment = wrap_top
                if row_type == "heading":
                    cell.fill = heading_fill
                    cell.font = Font(bold=True)
                elif row_type == "decision":
                    cell.fill = decision_fill
                elif row_type == "branch":
                    cell.fill = branch_ja_fill if row.get("branch") == "JA" else branch_nein_fill
                elif row_type == "hint":
                    cell.fill = hint_fill
            row_idx += 1

    widths = [12, 12, 55, 26, 30, 24, 22]
    for ci, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(ci)].width = w
    last_col = get_column_letter(len(headers))
    ws.auto_filter.ref = f"A1:{last_col}{max(1, row_idx - 1)}"
    ws.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── PDF-Export ──────────────────────────────────────────────────────────────────

def _esc(text: str | None) -> str:
    """Escaped HTML-Sonderzeichen fuer ReportLab-Paragraphen."""
    if not text:
        return ""
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def export_pdf(
    template: ChecklistTemplate,
    nodes: list[ChecklistTemplateNode],
    answer_sets: list[ChecklistAnswerSet],
    mode: str = "blank",
) -> bytes:
    """Erzeugt die ausfuellbare Pruefcheckliste als PDF (Bytes, reportlab).

    Strukturierte Tabelle analog zur DOCX-/XLSX-Ausgabe; ankreuzbare Optionen
    werden als ☐-Symbole, Freitextfelder als Linien dargestellt."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        Flowable, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    # ── Custom Flowables: echte AcroForm-Felder ───────────────────────────────
    class _AcroCheckbox(Flowable):
        """Anklickbare AcroForm-Checkbox (/Btn) als Flowable fester Groesse."""

        def __init__(self, name: str, size: float = 10):
            super().__init__()
            self._name = name
            self._size = size

        def wrap(self, *_args):
            self.width = self._size
            self.height = self._size
            return self.width, self.height

        def draw(self):
            self.canv.acroForm.checkbox(
                name=self._name, x=0, y=0, size=self._size,
                borderWidth=0.5, checked=False, relative=True,
            )

    class _AcroTextField(Flowable):
        """Ausfuellbares AcroForm-Textfeld (/Tx) als Flowable."""

        def __init__(self, name: str, width: float, height: float):
            super().__init__()
            self._name = name
            self.width = width
            self.height = height

        def wrap(self, *_args):
            return self.width, self.height

        def draw(self):
            self.canv.acroForm.textfield(
                name=self._name, x=0, y=0, width=self.width, height=self.height,
                borderWidth=0.5, fontSize=8, relative=True,
            )

    # Eindeutige Feldnamen ueber einen laufenden Zaehler.
    field_counter = {"n": 0}

    def _next(prefix: str) -> str:
        field_counter["n"] += 1
        return f"{prefix}_{field_counter['n']}"

    kopf, chapters = _prepare(template, nodes, answer_sets, mode)

    output = io.BytesIO()
    page_width, page_height = A4
    checklist_title = template.title or "Pruefcheckliste"

    def add_header_footer(canvas, doc_template):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#E0E0E0"))
        canvas.setLineWidth(0.5)
        canvas.line(15 * mm, page_height - 14 * mm, page_width - 15 * mm, page_height - 14 * mm)
        canvas.setFont("Helvetica-Bold", 9)
        canvas.setFillColor(colors.HexColor("#0E7C5C"))
        canvas.drawString(15 * mm, page_height - 12 * mm, "Pruefcheckliste")
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawRightString(page_width - 15 * mm, page_height - 12 * mm, checklist_title[:70])
        canvas.line(15 * mm, 12 * mm, page_width - 15 * mm, 12 * mm)
        canvas.drawString(15 * mm, 8 * mm, "Auditworkshop · KOM-Musterchecklist")
        canvas.drawRightString(page_width - 15 * mm, 8 * mm, f"Seite {doc_template.page}")
        canvas.restoreState()

    doc_template = SimpleDocTemplate(
        output, pagesize=A4,
        rightMargin=15 * mm, leftMargin=15 * mm,
        topMargin=20 * mm, bottomMargin=18 * mm,
    )

    primary = colors.HexColor("#0E7C5C")
    gray_light = colors.HexColor("#F1F5F9")
    gray_dark = colors.HexColor("#1F2937")
    grid = colors.HexColor("#D0D7DE")

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("CoverTitle", parent=styles["Heading1"], fontSize=20,
                              textColor=colors.white, alignment=TA_CENTER))
    styles.add(ParagraphStyle("Cell", parent=styles["Normal"], fontSize=8, leading=10,
                              textColor=gray_dark))
    styles.add(ParagraphStyle("CellBold", parent=styles["Normal"], fontSize=8, leading=10,
                              textColor=gray_dark, fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle("ChapterHead", parent=styles["Heading2"], fontSize=13,
                              textColor=primary, spaceBefore=6 * mm, spaceAfter=3 * mm))

    story: list[Any] = []

    # ── Deckblatt ──
    cover = Table([[Paragraph("Pruefcheckliste", styles["CoverTitle"])]], colWidths=[180 * mm])
    cover.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), primary),
        ("TOPPADDING", (0, 0), (-1, -1), 14),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 14),
    ]))
    story.append(cover)
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph(f"<b>{_esc(checklist_title)}</b>",
                          ParagraphStyle("T", parent=styles["Normal"], fontSize=14,
                                         textColor=gray_dark, spaceAfter=4 * mm)))
    if template.description:
        story.append(Paragraph(f"<i>{_esc(template.description)}</i>", styles["Cell"]))
        story.append(Spacer(1, 3 * mm))

    if kopf:
        info_data = [
            [Paragraph(f"<b>{_esc(label)}</b>", styles["Cell"]),
             Paragraph(_esc(value) if mode == "filled" and value else (_esc(value) or "&nbsp;"), styles["Cell"])]
            for label, value in kopf
        ]
        info_table = Table(info_data, colWidths=[45 * mm, 135 * mm])
        info_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), gray_light),
            ("GRID", (0, 0), (-1, -1), 0.5, grid),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(info_table)

    story.append(Spacer(1, 3 * mm))
    story.append(Paragraph(
        f"Erzeugt am {datetime.now().strftime('%d.%m.%Y %H:%M')} · Modus: {mode}",
        styles["Cell"],
    ))
    story.append(PageBreak())

    # ── Inhalt kapitelweise ──
    headers = ["Nr.", "Rechtsgrundlage", "Frage", "Antwort", "Bemerkung", "Belege"]
    col_widths = [11 * mm, 24 * mm, 56 * mm, 28 * mm, 32 * mm, 29 * mm]

    for chapter_title, rows in chapters:
        if chapter_title:
            story.append(Paragraph(_esc(chapter_title), styles["ChapterHead"]))
        if not rows:
            continue

        table_data: list[list[Any]] = [
            [Paragraph(f"<b>{_esc(h)}</b>",
                       ParagraphStyle("H", parent=styles["Cell"], textColor=colors.white))
             for h in headers]
        ]
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), primary),
            ("GRID", (0, 0), (-1, -1), 0.5, grid),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]

        for ri, row in enumerate(rows, start=1):
            row_type = row["row_type"]
            indent = "&nbsp;" * (4 * row.get("indent", 0))

            if row_type in ("heading", "branch", "hint"):
                label = row["question"]
                if row_type == "hint":
                    label = f"Hinweis: {label}"
                cell_para = Paragraph(
                    f"{indent}<b>{_esc(label)}</b>" if row_type in ("heading", "branch")
                    else f"{indent}<i>{_esc(label)}</i>",
                    styles["Cell"],
                )
                table_data.append([cell_para, "", "", "", "", ""])
                style_cmds.append(("SPAN", (0, ri), (-1, ri)))
                bg = {
                    "heading": colors.HexColor(f"#{HEADING_BG}"),
                    "hint": colors.HexColor(f"#{HINT_BG}"),
                    "branch": colors.HexColor(f"#{BRANCH_JA_BG}") if row.get("branch") == "JA"
                    else colors.HexColor(f"#{BRANCH_NEIN_BG}"),
                }[row_type]
                style_cmds.append(("BACKGROUND", (0, ri), (-1, ri), bg))
                continue

            # Antwort-Zelle: echte anklickbare AcroForm-Checkboxen je Option,
            # bei Freitext/Betrag/Datum ein AcroForm-Textfeld.
            answer_col_w = col_widths[3]
            if row["options"]:
                opt_rows: list[list[Any]] = []
                for opt in row["options"]:
                    opt_rows.append([
                        _AcroCheckbox(_next("ans"), size=10),
                        Paragraph(_esc(opt), styles["Cell"]),
                    ])
                ans_cell: Any = Table(opt_rows, colWidths=[5 * mm, answer_col_w - 5 * mm])
                ans_cell.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 1),
                    ("TOPPADDING", (0, 0), (-1, -1), 1),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ]))
            elif row.get("freitext_hint"):
                ans_cell = _AcroTextField(_next("ans"), width=answer_col_w - 3 * mm, height=5 * mm)
            else:
                ans_cell = _AcroTextField(_next("ans"), width=answer_col_w - 3 * mm, height=5 * mm)

            # Bemerkungs-Zelle: ausfuellbares AcroForm-Textfeld.
            remark_cell = _AcroTextField(_next("rmk"), width=col_widths[4] - 3 * mm, height=5 * mm)

            q_text = f"{indent}{_esc(row['question'])}"
            q_style = styles["CellBold"] if row_type == "decision" else styles["Cell"]
            table_data.append([
                Paragraph(_esc(row["nr"]), styles["Cell"]),
                Paragraph(_esc(row["legal"]), styles["Cell"]),
                Paragraph(q_text, q_style),
                ans_cell,
                remark_cell,
                Paragraph(_esc(row.get("documents", "")).replace("\n", "<br/>"), styles["Cell"]),
            ])
            if row_type == "decision":
                style_cmds.append(("BACKGROUND", (0, ri), (-1, ri), colors.HexColor(f"#{DECISION_BG}")))

        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle(style_cmds))
        story.append(table)
        story.append(Spacer(1, 4 * mm))

    doc_template.build(story, onFirstPage=add_header_footer, onLaterPages=add_header_footer)
    return output.getvalue()


# ── Diskussions-Export (Diskussionsprotokoll) ──────────────────────────────────

# Anzeige-Labels fuer den Team-Workflow-Status eines Knotens.
NODE_STATUS_LABELS = {
    "pending": "Offen",
    "in_progress": "In Bearbeitung",
    "resolved": "Erledigt",
}


def _format_dt(dt: datetime | None) -> str:
    """Formatiert einen Zeitstempel deutsch (oder leerer String)."""
    if not dt:
        return ""
    try:
        return dt.strftime("%d.%m.%Y %H:%M")
    except Exception:  # noqa: BLE001 — defensiv bei untypischen Werten
        return str(dt)


def _comment_meta(comment, names: dict[str, str | None]) -> str:
    """Baut die Kopfzeile eines Diskussionsbeitrags: Autor · Datum (· bearbeitet)."""
    author = names.get(getattr(comment, "author_id", None) or "") or "Unbekannt"
    created = _format_dt(getattr(comment, "created_at", None))
    meta = f"{author}"
    if created:
        meta += f" · {created}"
    if getattr(comment, "edited_at", None):
        meta += " · bearbeitet"
    return meta


def _threaded_comments(comments: list) -> list[tuple[Any, int]]:
    """Ordnet die (zeitlich sortierten) Kommentare als Thread:
    Wurzelkommentare auf Ebene 0, direkte Antworten auf Ebene 1.

    Rueckgabe: Liste von (comment, depth) in Anzeige-Reihenfolge.
    """
    by_parent: dict[str | None, list] = {}
    ids = {getattr(c, "id", None) for c in comments}
    for c in comments:
        parent = getattr(c, "parent_comment_id", None)
        if parent not in ids:
            parent = None
        by_parent.setdefault(parent, []).append(c)

    ordered: list[tuple[Any, int]] = []

    def _emit(parent_id, depth: int) -> None:
        for c in by_parent.get(parent_id, []):
            ordered.append((c, depth))
            # Genau eine Antwort-Ebene; tiefere Antworten haengen unter der Wurzel.
            _emit(getattr(c, "id", None), min(depth + 1, 1))

    _emit(None, 0)
    return ordered


def _discussion_nodes_in_order(
    nodes: list[ChecklistTemplateNode],
    comments_by_node: dict[str, list],
) -> list[tuple[str, ChecklistTemplateNode]]:
    """Liefert die Knoten MIT Kommentaren in Baum-Reihenfolge inkl. Nummerierung.

    Nutzt _build_node_tree fuer Reihenfolge und hierarchische Nummern. Knoten
    ohne Kommentare werden uebersprungen (aber ihre Nummerierungsposition bleibt
    in der Baum-Logik erhalten)."""
    by_id, roots = _build_node_tree(nodes)
    result: list[tuple[str, ChecklistTemplateNode]] = []

    def _walk(node_id: str, prefix: list[int]) -> None:
        entry = by_id.get(node_id)
        if not entry:
            return
        node: ChecklistTemplateNode = entry["node"]
        node_type = (node.node_type or "QUESTION").upper()
        children = entry["children"]
        nr = ".".join(str(p) for p in prefix) if prefix else ""
        if node_type != "HEADING" and comments_by_node.get(node_id):
            result.append((nr, node))
        counter = 0
        for child_id in children:
            counter += 1
            _walk(child_id, prefix + [counter])

    counter = 0
    for root_id in roots:
        counter += 1
        _walk(root_id, [counter])
    return result


def export_discussion_docx(
    template: ChecklistTemplate,
    nodes: list[ChecklistTemplateNode],
    comments: list,
    names: dict[str, str | None],
) -> bytes:
    """Erzeugt das Diskussionsprotokoll als DOCX (Bytes).

    Nur Knoten mit Kommentaren, in Baum-Reihenfolge. Je Knoten: Nummer, Titel,
    Status-Label, Rechtsgrundlage; darunter die Kommentare threaded (Antworten
    eingerueckt). Soft-geloeschte Beitraege erscheinen als ``[gelöscht]``."""
    # Kommentare je Knoten gruppieren (Eingangsreihenfolge = created_at asc).
    comments_by_node: dict[str, list] = {}
    for c in comments:
        comments_by_node.setdefault(c.node_id, []).append(c)

    ordered_nodes = _discussion_nodes_in_order(nodes, comments_by_node)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ── Kopfblock ──
    header_table = doc.add_table(rows=1, cols=1)
    header_cell = header_table.rows[0].cells[0]
    _set_cell_shading(header_cell, HEADER_BG)
    title_para = header_cell.paragraphs[0]
    title_run = title_para.add_run("Diskussionsprotokoll")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(255, 255, 255)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(14)
    sub_para = header_cell.add_paragraph()
    sub_run = sub_para.add_run(template.title or "")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(255, 255, 255)
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()
    gen = doc.add_paragraph()
    gen_run = gen.add_run(f"Erzeugt am {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    gen_run.font.size = Pt(8)
    gen_run.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()

    if not ordered_nodes:
        info = doc.add_paragraph()
        info.add_run("Keine Diskussionsbeiträge vorhanden.").italic = True
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    for nr, node in ordered_nodes:
        title = _node_title(node)
        status_label = NODE_STATUS_LABELS.get((node.status or "").lower(), "")
        heading_text = f"{nr} {title}".strip()
        h = doc.add_heading(heading_text or "Knoten", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(14, 124, 92)

        meta_para = doc.add_paragraph()
        meta_bits = []
        if status_label:
            meta_bits.append(f"Status: {status_label}")
        if (node.legal_reference or "").strip():
            meta_bits.append(f"Rechtsgrundlage: {node.legal_reference.strip()}")
        if meta_bits:
            mr = meta_para.add_run("   ·   ".join(meta_bits))
            mr.font.size = Pt(8)
            mr.font.color.rgb = RGBColor(107, 114, 128)

        for comment, depth in _threaded_comments(comments_by_node.get(node.id, [])):
            indent = Cm(0.6 * depth)
            deleted = getattr(comment, "deleted_at", None) is not None

            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.left_indent = indent
            meta_p.paragraph_format.space_after = Pt(1)
            mrun = meta_p.add_run(_comment_meta(comment, names))
            mrun.bold = True
            mrun.font.size = Pt(9)

            msg_p = doc.add_paragraph()
            msg_p.paragraph_format.left_indent = indent
            msg_p.paragraph_format.space_after = Pt(6)
            text = "[gelöscht]" if deleted else (comment.message or "")
            trun = msg_p.add_run(text)
            trun.font.size = Pt(10)
            if deleted:
                trun.italic = True
                trun.font.color.rgb = RGBColor(150, 150, 150)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_discussion_pdf(
    template: ChecklistTemplate,
    nodes: list[ChecklistTemplateNode],
    comments: list,
    names: dict[str, str | None],
) -> bytes:
    """Erzeugt das Diskussionsprotokoll als PDF (Bytes, reportlab)."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    comments_by_node: dict[str, list] = {}
    for c in comments:
        comments_by_node.setdefault(c.node_id, []).append(c)
    ordered_nodes = _discussion_nodes_in_order(nodes, comments_by_node)

    output = io.BytesIO()
    page_width, page_height = A4
    proto_title = template.title or "Diskussionsprotokoll"

    def add_header_footer(canvas, doc_template):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#E0E0E0"))
        canvas.setLineWidth(0.5)
        canvas.line(15 * mm, page_height - 14 * mm, page_width - 15 * mm, page_height - 14 * mm)
        canvas.setFont("Helvetica-Bold", 9)
        canvas.setFillColor(colors.HexColor("#0E7C5C"))
        canvas.drawString(15 * mm, page_height - 12 * mm, "Diskussionsprotokoll")
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawRightString(page_width - 15 * mm, page_height - 12 * mm, proto_title[:70])
        canvas.line(15 * mm, 12 * mm, page_width - 15 * mm, 12 * mm)
        canvas.drawString(15 * mm, 8 * mm, "Auditworkshop · KOM-Checklisten-Designer")
        canvas.drawRightString(page_width - 15 * mm, 8 * mm, f"Seite {doc_template.page}")
        canvas.restoreState()

    doc_template = SimpleDocTemplate(
        output, pagesize=A4,
        rightMargin=15 * mm, leftMargin=15 * mm,
        topMargin=20 * mm, bottomMargin=18 * mm,
    )

    primary = colors.HexColor("#0E7C5C")
    gray_dark = colors.HexColor("#1F2937")
    gray_mid = colors.HexColor("#6B7280")

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("DiscTitle", parent=styles["Heading1"], fontSize=20,
                              textColor=colors.white, alignment=TA_CENTER))
    styles.add(ParagraphStyle("Node", parent=styles["Heading2"], fontSize=13,
                              textColor=primary, spaceBefore=6 * mm, spaceAfter=1 * mm))
    styles.add(ParagraphStyle("NodeMeta", parent=styles["Normal"], fontSize=8,
                              textColor=gray_mid, spaceAfter=3 * mm))
    styles.add(ParagraphStyle("CMeta", parent=styles["Normal"], fontSize=9,
                              textColor=gray_dark, fontName="Helvetica-Bold", spaceAfter=1))
    styles.add(ParagraphStyle("CMsg", parent=styles["Normal"], fontSize=10,
                              textColor=gray_dark, leading=13, spaceAfter=4))
    styles.add(ParagraphStyle("CMsgDel", parent=styles["Normal"], fontSize=10,
                              textColor=gray_mid, leading=13, spaceAfter=4))

    story: list[Any] = []
    cover = Table([[Paragraph("Diskussionsprotokoll", styles["DiscTitle"])]], colWidths=[180 * mm])
    cover.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), primary),
        ("TOPPADDING", (0, 0), (-1, -1), 14),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 14),
    ]))
    story.append(cover)
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph(f"<b>{_esc(proto_title)}</b>",
                          ParagraphStyle("PT", parent=styles["Normal"], fontSize=14,
                                         textColor=gray_dark, spaceAfter=2 * mm)))
    story.append(Paragraph(
        f"Erzeugt am {datetime.now().strftime('%d.%m.%Y %H:%M')}", styles["NodeMeta"]))

    if not ordered_nodes:
        story.append(Spacer(1, 4 * mm))
        story.append(Paragraph("<i>Keine Diskussionsbeiträge vorhanden.</i>", styles["CMsg"]))
        doc_template.build(story, onFirstPage=add_header_footer, onLaterPages=add_header_footer)
        return output.getvalue()

    for nr, node in ordered_nodes:
        title = _node_title(node)
        story.append(Paragraph(f"{_esc(nr)} {_esc(title)}".strip(), styles["Node"]))

        status_label = NODE_STATUS_LABELS.get((node.status or "").lower(), "")
        meta_bits = []
        if status_label:
            meta_bits.append(f"Status: {_esc(status_label)}")
        if (node.legal_reference or "").strip():
            meta_bits.append(f"Rechtsgrundlage: {_esc(node.legal_reference.strip())}")
        if meta_bits:
            story.append(Paragraph("   ·   ".join(meta_bits), styles["NodeMeta"]))

        for comment, depth in _threaded_comments(comments_by_node.get(node.id, [])):
            deleted = getattr(comment, "deleted_at", None) is not None
            indent = "&nbsp;" * (6 * depth)
            story.append(Paragraph(f"{indent}{_esc(_comment_meta(comment, names))}", styles["CMeta"]))
            if deleted:
                story.append(Paragraph(f"{indent}<i>[gelöscht]</i>", styles["CMsgDel"]))
            else:
                msg = _esc(comment.message or "").replace("\n", "<br/>")
                story.append(Paragraph(f"{indent}{msg}", styles["CMsg"]))

    doc_template.build(story, onFirstPage=add_header_footer, onLaterPages=add_header_footer)
    return output.getvalue()
