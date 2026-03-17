"""
flowworkshop · file_parser.py
Universeller Datei-Parser: PDF, XLSX, XLS, XLSM, DOCX, HTML, RTF.

Nutzt pdf_parser.py fuer PDFs (3-Stufen-Fallback mit OCR).
"""
from __future__ import annotations
import io
import logging
from pathlib import Path

log = logging.getLogger(__name__)

# Erlaubte Dateiendungen
ALLOWED_EXTENSIONS = {
    ".pdf", ".xlsx", ".xls", ".xlsm",
    ".docx", ".docm", ".html", ".htm", ".rtf", ".txt",
}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB


def extract(file_bytes: bytes, filename: str) -> dict:
    """
    Extrahiert Text aus einer Datei.

    Returns:
        {
          "text": str,
          "method": str,       — z.B. "pdfplumber", "openpyxl", "python-docx", "html", "rtf"
          "pages": int,        — Seitenanzahl (0 bei Nicht-PDF)
          "char_count": int,
          "warnings": list[str]
        }
    """
    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError("Datei zu gross (max. 50 MB).")

    ext = Path(filename).suffix.lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Dateityp '{ext}' nicht unterstuetzt. "
            f"Erlaubt: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    if ext == ".pdf":
        from services.pdf_parser import extract as pdf_extract
        return pdf_extract(file_bytes, filename)

    if ext in (".xlsx", ".xls", ".xlsm"):
        return _extract_excel(file_bytes, filename, ext)

    if ext in (".docx", ".docm"):
        return _extract_docx(file_bytes, filename)

    if ext in (".html", ".htm"):
        return _extract_html(file_bytes, filename)

    if ext == ".rtf":
        return _extract_rtf(file_bytes, filename)

    if ext == ".txt":
        return _extract_text(file_bytes, filename)

    raise ValueError(f"Kein Parser fuer '{ext}' implementiert.")


def _extract_excel(file_bytes: bytes, filename: str, ext: str) -> dict:
    """
    Excel-Dateien (XLSX, XLS, XLSM) → getaggter Text.

    Jede Datenzeile wird als Key-Value-Paar mit Spaltenheader formatiert,
    damit jeder Chunk selbsterklaerend ist und vom LLM interpretiert werden kann.

    Format pro Datensatz:
        [Blatt: Name | Zeile 7]
        Spalte1: Wert1
        Spalte2: Wert2
        ...
    """
    warnings: list[str] = []
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []
        sheet_count = len(wb.sheetnames)

        for ws in wb.worksheets:
            # Header-Zeile finden (erste Zeile mit mindestens 2 nicht-leeren Zellen)
            headers: list[str] = []
            header_row_idx = 0
            data_rows: list[tuple[int, list]] = []

            for i, row in enumerate(ws.iter_rows(values_only=True)):
                cells = [c for c in row]
                non_empty = [str(c).strip() for c in cells if c is not None and str(c).strip()]

                if not headers and len(non_empty) >= 2:
                    # Das ist vermutlich die Header-Zeile
                    headers = [str(c).strip() if c is not None else f"Spalte_{j+1}"
                               for j, c in enumerate(cells)]
                    header_row_idx = i
                elif headers:
                    # Datenzeile — nur wenn mindestens ein Wert vorhanden
                    if non_empty:
                        data_rows.append((i + 1, list(cells)))

            if not headers or not data_rows:
                # Fallback: kein erkennbarer Header → altes Tab-Format
                rows_text = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        rows_text.append("\t".join(cells))
                if rows_text:
                    parts.append(f"=== Blatt: {ws.title} ===\n" + "\n".join(rows_text))
                continue

            # Getaggte Datenzeilen erstellen
            for row_num, cells in data_rows:
                tagged_fields = []
                for j, val in enumerate(cells):
                    if j >= len(headers):
                        break
                    header = headers[j]
                    if not header or header.startswith("Spalte_"):
                        continue
                    cell_val = str(val).strip() if val is not None else ""
                    if not cell_val or cell_val == "None":
                        continue
                    tagged_fields.append(f"  {header}: {cell_val}")

                if tagged_fields:
                    record = f"[Blatt: {ws.title} | Zeile {row_num}]\n" + "\n".join(tagged_fields)
                    parts.append(record)

        wb.close()
        text = "\n\n".join(parts).strip()
        return {
            "text": text, "method": "openpyxl-tagged",
            "pages": sheet_count, "char_count": len(text), "warnings": warnings,
        }
    except Exception as e:
        warnings.append(f"openpyxl: {e}")
        # Fallback mit pandas
        try:
            import pandas as pd
            engine = "openpyxl" if ext in (".xlsx", ".xlsm") else "xlrd"
            sheets = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, engine=engine)
            parts = []
            for name, df in sheets.items():
                for idx, row in df.iterrows():
                    fields = [f"  {col}: {val}" for col, val in row.items()
                              if pd.notna(val) and str(val).strip()]
                    if fields:
                        parts.append(f"[Blatt: {name} | Zeile {idx + 2}]\n" + "\n".join(fields))
            text = "\n\n".join(parts).strip()
            return {
                "text": text, "method": "pandas-tagged",
                "pages": len(sheets), "char_count": len(text), "warnings": warnings,
            }
        except Exception as e2:
            warnings.append(f"pandas: {e2}")
    return {"text": "", "method": "failed", "pages": 0, "char_count": 0, "warnings": warnings}


def _extract_docx(file_bytes: bytes, filename: str) -> dict:
    """DOCX/DOCM → Absatztext."""
    warnings: list[str] = []
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Tabellen auch extrahieren
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    paragraphs.append("\t".join(cells))
        text = "\n".join(paragraphs).strip()
        return {
            "text": text, "method": "python-docx",
            "pages": 0, "char_count": len(text), "warnings": warnings,
        }
    except Exception as e:
        warnings.append(f"python-docx: {e}")
    return {"text": "", "method": "failed", "pages": 0, "char_count": 0, "warnings": warnings}


def _extract_html(file_bytes: bytes, filename: str) -> dict:
    """HTML → Klartext."""
    warnings: list[str] = []
    try:
        # Encoding erraten
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                html = file_bytes.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        else:
            html = file_bytes.decode("utf-8", errors="replace")

        # Einfache Tag-Entfernung (keine externe Dependency noetig)
        import re
        text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return {
            "text": text, "method": "html-strip",
            "pages": 0, "char_count": len(text), "warnings": warnings,
        }
    except Exception as e:
        warnings.append(f"html: {e}")
    return {"text": "", "method": "failed", "pages": 0, "char_count": 0, "warnings": warnings}


def _extract_rtf(file_bytes: bytes, filename: str) -> dict:
    """RTF → Klartext (einfache Entfernung der RTF-Steuerzeichen)."""
    warnings: list[str] = []
    try:
        import re
        raw = file_bytes.decode("latin-1", errors="replace")
        # RTF-Gruppen und Steuerzeichen entfernen
        text = re.sub(r"\\[a-z]+\d*\s?", " ", raw)
        text = re.sub(r"[{}]", "", text)
        text = re.sub(r"\s+", " ", text).strip()
        if text and len(text) > 20:
            return {
                "text": text, "method": "rtf-strip",
                "pages": 0, "char_count": len(text), "warnings": warnings,
            }
        warnings.append("RTF: zu wenig Text nach Parsing")
    except Exception as e:
        warnings.append(f"rtf: {e}")
    return {"text": "", "method": "failed", "pages": 0, "char_count": 0, "warnings": warnings}


def _extract_text(file_bytes: bytes, filename: str) -> dict:
    """Reine Textdatei."""
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            text = file_bytes.decode(enc).strip()
            return {
                "text": text, "method": "text",
                "pages": 0, "char_count": len(text), "warnings": [],
            }
        except UnicodeDecodeError:
            continue
    text = file_bytes.decode("utf-8", errors="replace").strip()
    return {
        "text": text, "method": "text",
        "pages": 0, "char_count": len(text), "warnings": ["Encoding-Fallback"],
    }
